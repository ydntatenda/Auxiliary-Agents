"""Document source ingestion.

Documents come in three shapes that cover most business artifacts:
PDF (born-digital or scanned), docx (Word), and images of forms or
screenshots. For PDFs we first try pypdf for embedded text; if a page
yields nothing, we treat the PDF as scanned and ask Gemini to read it
as images. docx uses python-docx. Standalone images go straight to
Gemini vision.

Model split is preserved: text and audio go to OpenAI elsewhere; images
and video go to Gemini. PDF rasterisation only happens when the embedded
text path comes up empty.
"""
from __future__ import annotations

import asyncio
from pathlib import Path

from app.config import get_settings
from app.services.gemini_client import get_gemini_client


IMAGE_EXTRACTION_PROMPT = (
    "You are reading a single image of a business document, form, or "
    "screenshot. Produce a clean plain-text rendering of every visible "
    "field label, value, heading, instruction, and table row. Preserve "
    "the visual order. Do not invent content you cannot see. If a field "
    "is empty, write the label followed by an empty value. Do not add "
    "commentary."
)

SCANNED_PDF_PROMPT = (
    "You are reading a scanned PDF page-by-page. For each page produce a "
    "clean plain-text rendering of every visible field label, value, "
    "heading, and table row, preserving order within the page. Begin "
    "each page with a line of the form '--- Page N ---'. Do not invent "
    "content you cannot see and do not add commentary."
)


async def ingest_document(raw_path: str) -> "IngestResult":
    from .skill import IngestResult  # local import avoids circular dep at module load

    path = Path(raw_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return await _ingest_pdf(path)
    if suffix == ".docx":
        return await asyncio.to_thread(_ingest_docx, path)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
        text = await _gemini_read_image(path)
        return IngestResult(
            assembled_text=text,
            meta={"ocr_used": True, "page_count": 1, "char_count": len(text)},
        )
    raise ValueError(f"unsupported document extension {suffix!r}")


async def _ingest_pdf(path: Path) -> "IngestResult":
    from .skill import IngestResult

    from pypdf import PdfReader  # noqa: PLC0415

    reader = await asyncio.to_thread(PdfReader, str(path))
    page_count = len(reader.pages)
    parts: list[str] = []
    for page in reader.pages:
        text = await asyncio.to_thread(page.extract_text) or ""
        parts.append(text.strip())

    combined = "\n\n".join(p for p in parts if p)
    if combined.strip():
        return IngestResult(
            assembled_text=combined,
            meta={"ocr_used": False, "page_count": page_count, "char_count": len(combined)},
        )

    text = await _gemini_read_scanned_pdf(path)
    return IngestResult(
        assembled_text=text,
        meta={"ocr_used": True, "page_count": page_count, "char_count": len(text)},
    )


def _ingest_docx(path: Path) -> "IngestResult":
    from .skill import IngestResult

    from docx import Document  # noqa: PLC0415

    document = Document(str(path))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    combined = "\n".join(paragraphs)
    return IngestResult(
        assembled_text=combined,
        meta={"ocr_used": False, "page_count": None, "char_count": len(combined)},
    )


async def _gemini_read_image(path: Path) -> str:
    client = get_gemini_client()
    settings = get_settings()
    uploaded = await asyncio.to_thread(client.files.upload, file=str(path))
    response = await asyncio.to_thread(
        client.models.generate_content,
        model=settings.gemini_video_model,
        contents=[uploaded, IMAGE_EXTRACTION_PROMPT],
    )
    return response.text or ""


async def _gemini_read_scanned_pdf(path: Path) -> str:
    """Send the whole PDF to Gemini for OCR-style read.

    Gemini 2.5 Pro accepts PDF uploads directly and treats each page as an
    image internally, which is exactly what we want here. We avoid the extra
    rasterisation hop (pdf2image / poppler) that would otherwise add a system
    dependency on the container.
    """
    client = get_gemini_client()
    settings = get_settings()
    uploaded = await asyncio.to_thread(client.files.upload, file=str(path))
    response = await asyncio.to_thread(
        client.models.generate_content,
        model=settings.gemini_video_model,
        contents=[uploaded, SCANNED_PDF_PROMPT],
    )
    return response.text or ""
